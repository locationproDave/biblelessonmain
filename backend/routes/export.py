# Export Routes
from fastapi import APIRouter, HTTPException, Header
from datetime import datetime, timezone
import json
import io
import base64
import logging
import re

from models.schemas import ExportRequest
from services.database import db

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/export", tags=["Export"])

def strip_emojis(text: str) -> str:
    """Remove emoji characters from text for PDF compatibility"""
    emoji_pattern = re.compile("["
        u"\U0001F600-\U0001F64F"
        u"\U0001F300-\U0001F5FF"
        u"\U0001F680-\U0001F6FF"
        u"\U0001F1E0-\U0001F1FF"
        u"\U00002702-\U000027B0"
        u"\U000024C2-\U0001F251"
        "]+", flags=re.UNICODE)
    return emoji_pattern.sub('', text).strip()

def serialize_doc(doc: dict) -> dict:
    if doc is None:
        return None
    result = {k: v for k, v in doc.items() if k != '_id'}
    for key, value in result.items():
        if isinstance(value, datetime):
            result[key] = value.isoformat()
    return result

async def get_current_user(token: str = None) -> dict:
    if not token:
        return None
    session = await db.sessions.find_one({"token": token})
    if not session:
        return None
    user = await db.users.find_one({"id": session["userId"]})
    return serialize_doc(user)

# Import PDF and DOCX libraries
try:
    from fpdf import FPDF
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

class BibleLessonPDF(FPDF):
    """Custom PDF class for Bible lessons"""
    def __init__(self):
        super().__init__()
        self.add_page()
        self.set_auto_page_break(auto=True, margin=15)
    
    def header(self):
        self.set_font('Helvetica', 'B', 10)
        self.set_text_color(30, 58, 95)
        self.cell(0, 10, 'Bible Lesson Planner', align='R')
        self.ln(15)
    
    def footer(self):
        self.set_y(-15)
        self.set_font('Helvetica', 'I', 8)
        self.set_text_color(107, 114, 128)
        self.cell(0, 10, f'Page {self.page_no()}', align='C')

def generate_pdf_bytes(lesson: dict, sections: list, materials: list) -> bytes:
    """Generate a PDF document from lesson data"""
    pdf = BibleLessonPDF()
    
    # Title
    pdf.set_font('Helvetica', 'B', 24)
    pdf.set_text_color(30, 58, 95)
    pdf.multi_cell(0, 10, strip_emojis(lesson.get('title', 'Untitled Lesson')))
    pdf.ln(5)
    
    # Meta info
    pdf.set_font('Helvetica', '', 10)
    pdf.set_text_color(107, 114, 128)
    meta_text = f"Passage: {lesson.get('passage', 'N/A')} | Age Group: {lesson.get('ageGroup', 'N/A')} | Duration: {lesson.get('duration', 'N/A')}"
    pdf.multi_cell(0, 6, meta_text)
    pdf.ln(10)
    
    # Memory Verse Box
    pdf.set_fill_color(254, 243, 199)
    pdf.set_draw_color(212, 160, 23)
    pdf.rect(10, pdf.get_y(), 190, 30, 'DF')
    pdf.set_y(pdf.get_y() + 5)
    pdf.set_font('Helvetica', 'B', 12)
    pdf.set_text_color(30, 58, 95)
    pdf.cell(0, 6, 'Memory Verse', ln=True)
    pdf.set_font('Helvetica', 'I', 11)
    pdf.set_text_color(55, 65, 81)
    verse_text = f'"{strip_emojis(lesson.get("memoryVerseText", ""))}"'
    pdf.multi_cell(180, 5, verse_text)
    pdf.set_font('Helvetica', 'B', 10)
    pdf.cell(0, 5, f"- {lesson.get('memoryVerseReference', '')}")
    pdf.ln(15)
    
    # Objectives
    pdf.set_font('Helvetica', 'B', 14)
    pdf.set_text_color(30, 58, 95)
    pdf.cell(0, 8, 'Learning Objectives', ln=True)
    pdf.set_draw_color(212, 160, 23)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(5)
    
    pdf.set_font('Helvetica', '', 11)
    pdf.set_text_color(55, 65, 81)
    for i, obj in enumerate(lesson.get('objectives', []), 1):
        pdf.multi_cell(0, 6, f"{i}. {strip_emojis(obj)}")
    pdf.ln(8)
    
    # Lesson Content
    pdf.set_font('Helvetica', 'B', 14)
    pdf.set_text_color(30, 58, 95)
    pdf.cell(0, 8, 'Lesson Content', ln=True)
    pdf.set_draw_color(212, 160, 23)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(5)
    
    for section in sections:
        if pdf.get_y() > 250:
            pdf.add_page()
        
        pdf.set_fill_color(245, 248, 252)
        pdf.set_font('Helvetica', 'B', 12)
        pdf.set_text_color(30, 58, 95)
        title = strip_emojis(f"{section.get('title', 'Section')}")
        pdf.cell(0, 8, title, fill=True, ln=True)
        
        pdf.set_font('Helvetica', 'I', 9)
        pdf.set_text_color(107, 114, 128)
        pdf.cell(0, 5, f"Duration: {section.get('duration', 'N/A')}", ln=True)
        
        pdf.set_font('Helvetica', '', 10)
        pdf.set_text_color(55, 65, 81)
        content = strip_emojis(section.get('content', ''))
        pdf.multi_cell(0, 5, content)
        pdf.ln(5)
    
    if materials:
        if pdf.get_y() > 220:
            pdf.add_page()
        
        pdf.set_font('Helvetica', 'B', 14)
        pdf.set_text_color(30, 58, 95)
        pdf.cell(0, 8, 'Materials Needed', ln=True)
        pdf.set_draw_color(212, 160, 23)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(5)
        
        pdf.set_fill_color(232, 245, 233)
        pdf.set_font('Helvetica', '', 10)
        pdf.set_text_color(55, 65, 81)
        
        for mat in materials:
            if isinstance(mat, dict):
                item_text = f"- {strip_emojis(mat.get('item', 'Item'))}"
                if mat.get('category'):
                    item_text += f" ({mat.get('category')})"
            else:
                item_text = f"- {strip_emojis(str(mat))}"
            pdf.multi_cell(0, 5, item_text)
    
    return pdf.output()

def generate_docx_bytes(lesson: dict, sections: list, materials: list) -> bytes:
    """Generate a Word document from lesson data"""
    doc = Document()
    
    # Title
    title = doc.add_heading(lesson.get('title', 'Untitled Lesson'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(30, 58, 95)
    
    # Meta info
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"Passage: {lesson.get('passage', 'N/A')} | Age Group: {lesson.get('ageGroup', 'N/A')} | Duration: {lesson.get('duration', 'N/A')}"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(107, 114, 128)
    
    # Memory Verse
    doc.add_heading('Memory Verse', level=1)
    verse_para = doc.add_paragraph()
    verse_run = verse_para.add_run(f'"{lesson.get("memoryVerseText", "")}"')
    verse_run.italic = True
    verse_run.font.size = Pt(12)
    ref_para = doc.add_paragraph()
    ref_run = ref_para.add_run(f"— {lesson.get('memoryVerseReference', '')}")
    ref_run.bold = True
    ref_run.font.size = Pt(11)
    
    # Objectives
    doc.add_heading('Learning Objectives', level=1)
    for i, obj in enumerate(lesson.get('objectives', []), 1):
        doc.add_paragraph(f"{i}. {obj}", style='List Number')
    
    # Lesson Content
    doc.add_heading('Lesson Content', level=1)
    for section in sections:
        section_title = f"{section.get('icon', '📖')} {section.get('title', 'Section')}"
        doc.add_heading(section_title, level=2)
        
        duration_para = doc.add_paragraph()
        dur_run = duration_para.add_run(f"Duration: {section.get('duration', 'N/A')}")
        dur_run.italic = True
        dur_run.font.size = Pt(9)
        dur_run.font.color.rgb = RGBColor(107, 114, 128)
        
        content = section.get('content', '')
        doc.add_paragraph(content)
    
    if materials:
        doc.add_heading('Materials Needed', level=1)
        for mat in materials:
            if isinstance(mat, dict):
                item_text = mat.get('item', 'Item')
                if mat.get('category'):
                    item_text += f" ({mat.get('category')})"
            else:
                item_text = str(mat)
            doc.add_paragraph(item_text, style='List Bullet')
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

@router.post("/lesson")
async def export_lesson(request: ExportRequest, authorization: str = Header(None)):
    """Export a lesson to PDF or Word format"""
    if not authorization:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    token = authorization.replace("Bearer ", "")
    user = await get_current_user(token)
    if not user:
        raise HTTPException(status_code=401, detail="Invalid token")
    
    if not PDF_AVAILABLE:
        raise HTTPException(status_code=500, detail="PDF generation libraries not available")
    
    lesson = await db.lessons.find_one({"id": request.lessonId})
    if not lesson:
        raise HTTPException(status_code=404, detail="Lesson not found")
    
    try:
        sections_raw = lesson.get("sectionsJson", "[]")
        materials_raw = lesson.get("materialsJson", "[]")
        
        if isinstance(sections_raw, str):
            sections = json.loads(sections_raw)
        else:
            sections = sections_raw
            
        if isinstance(materials_raw, str):
            materials = json.loads(materials_raw)
        else:
            materials = materials_raw
        
        if request.format == "pdf":
            file_bytes = generate_pdf_bytes(lesson, sections, materials)
            content_type = "application/pdf"
            filename = f"{lesson.get('title', 'lesson').replace(' ', '_')}.pdf"
        elif request.format == "docx":
            file_bytes = generate_docx_bytes(lesson, sections, materials)
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"{lesson.get('title', 'lesson').replace(' ', '_')}.docx"
        else:
            raise HTTPException(status_code=400, detail="Invalid format. Use 'pdf' or 'docx'")
        
        file_base64 = base64.b64encode(file_bytes).decode('utf-8')
        
        return {
            "success": True,
            "lessonId": request.lessonId,
            "format": request.format,
            "filename": filename,
            "contentType": content_type,
            "fileData": file_base64,
            "metadata": {
                "title": lesson.get('title'),
                "passage": lesson.get('passage'),
                "ageGroup": lesson.get('ageGroup'),
                "duration": lesson.get('duration'),
                "createdAt": lesson.get('createdAt')
            }
        }
        
    except Exception as e:
        logger.error(f"Export error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")

@router.get("/print/{lesson_id}")
async def get_print_friendly_lesson(lesson_id: str):
    """Get print-friendly HTML version of a lesson"""
    lesson = await db.lessons.find_one({"id": lesson_id})
    if not lesson:
        raise HTTPException(status_code=404, detail="Lesson not found")
    
    try:
        sections_raw = lesson.get("sectionsJson", "[]")
        materials_raw = lesson.get("materialsJson", "[]")
        
        if isinstance(sections_raw, str):
            sections = json.loads(sections_raw)
        else:
            sections = sections_raw
            
        if isinstance(materials_raw, str):
            materials = json.loads(materials_raw)
        else:
            materials = materials_raw
        
        # Generate sections HTML
        sections_html = ""
        for section in sections:
            if isinstance(section, dict):
                sections_html += f"""
                <div class="section">
                    <h3>{section.get('icon', '📖')} {section.get('title', 'Section')}</h3>
                    <p class="duration">Duration: {section.get('duration', 'N/A')}</p>
                    <p>{section.get('content', '')}</p>
                </div>
                """
        
        # Generate materials HTML
        materials_html = ""
        for mat in materials:
            if isinstance(mat, dict):
                item = mat.get('item', 'Item')
                category = mat.get('category', '')
                materials_html += f"<li>{item} {f'({category})' if category else ''}</li>"
            else:
                materials_html += f"<li>{mat}</li>"
        
        # Generate objectives HTML
        objectives_html = ""
        for obj in lesson.get('objectives', []):
            objectives_html += f"<li>{obj}</li>"
        
        # Cross references
        cross_refs_json = lesson.get("crossReferencesJson", "[]")
        if isinstance(cross_refs_json, str):
            try:
                cross_refs = json.loads(cross_refs_json)
            except:
                cross_refs = []
        else:
            cross_refs = cross_refs_json
        
        cross_refs_html = ""
        for ref in cross_refs:
            if isinstance(ref, dict):
                cross_refs_html += f"<li><strong>{ref.get('reference', '')}</strong>: {ref.get('text', '')}</li>"
            else:
                cross_refs_html += f"<li>{ref}</li>"
        
        return {
            "lessonId": lesson_id,
            "title": lesson.get('title', 'Untitled Lesson'),
            "passage": lesson.get('passage', 'N/A'),
            "ageGroup": lesson.get('ageGroup', 'N/A'),
            "duration": lesson.get('duration', 'N/A'),
            "theme": lesson.get('theme', 'N/A'),
            "format": lesson.get('format', 'N/A'),
            "memoryVerse": {
                "text": lesson.get('memoryVerseText', ''),
                "reference": lesson.get('memoryVerseReference', '')
            },
            "objectives": lesson.get('objectives', []),
            "sections": sections,
            "materials": materials,
            "crossReferences": cross_refs,
            "description": lesson.get('description', ''),
            "html": {
                "sectionsHtml": sections_html,
                "materialsHtml": materials_html,
                "objectivesHtml": objectives_html,
                "crossRefsHtml": cross_refs_html
            }
        }
    except Exception as e:
        logger.error(f"Print export error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate print view: {str(e)}")
